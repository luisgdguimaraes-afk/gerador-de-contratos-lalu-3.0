"""
Serviço para extração de texto e identificação de placeholders em DOCX
"""
import re
from docx import Document
from typing import List, Dict, Tuple


class DocumentParser:
    """Extrai texto e identifica campos editáveis em documentos DOCX"""
    
    # Novo padrão: detecta {{CAMPO}}
    PLACEHOLDER_PATTERN = r'\{\{([A-Z0-9_]+)\}\}'
    
    def extract_text(self, docx_path: str) -> str:
        """
        Extrai todo o texto do documento DOCX incluindo parágrafos e tabelas
        """
        doc = Document(docx_path)
        full_text = []
        
        # Extrair de parágrafos
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # Extrair de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            full_text.append(text)
        
        return '\n'.join(full_text)
    
    def extract_text_with_structure(self, docx_path: str) -> List[Dict]:
        """
        Extrai texto mantendo estrutura (parágrafos, runs) para preservar formatação
        Retorna lista de parágrafos com informações de formatação
        """
        doc = Document(docx_path)
        structured_text = []
        
        for para in doc.paragraphs:
            para_data = {
                'text': para.text,
                'runs': []
            }
            
            for run in para.runs:
                run_data = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline
                }
                para_data['runs'].append(run_data)
            
            structured_text.append(para_data)
        
        return structured_text
    
    def find_placeholders(self, text: str) -> List[Dict]:
        """
        Encontra todos os placeholders no formato {{CAMPO}}
        Retorna lista de dicts com field_id, original_text, posição inicial e final
        """
        placeholders = []
        
        for match in re.finditer(self.PLACEHOLDER_PATTERN, text):
            placeholders.append({
                "field_id": match.group(1),  # Nome do campo sem {{ }}
                "original_text": match.group(0),  # Texto completo {{CAMPO}}
                "start": match.start(),
                "end": match.end()
            })
        
        return placeholders
    
    def validate_placeholders(self, found_placeholders: List[Dict], 
                               expected_fields: List[str]) -> Dict:
        """
        Valida se os placeholders encontrados correspondem aos esperados
        """
        found_ids = set([p["field_id"] for p in found_placeholders])
        expected_ids = set(expected_fields)
        
        return {
            "valid": found_ids == expected_ids,
            "missing": list(expected_ids - found_ids),
            "extra": list(found_ids - expected_ids),
            "found": list(found_ids)
        }
    
    def get_context_around_placeholder(self, text: str, start: int, end: int, 
                                      context_chars: int = 100) -> str:
        """
        Extrai contexto ao redor de um placeholder
        """
        context_start = max(0, start - context_chars)
        context_end = min(len(text), end + context_chars)
        
        return text[context_start:context_end]
