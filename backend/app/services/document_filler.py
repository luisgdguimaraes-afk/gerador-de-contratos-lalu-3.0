"""
Serviço para preencher documentos DOCX com dados do formulário
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from typing import Dict, Any, Optional
import re
from app.services.document_storage import DocumentStorage
from app.services.field_validator import FieldValidator
from app.services.contract_schema import ROTA_DO_SOL_SCHEMA, FieldType


class DocumentFiller:
    """Preenche campos em documentos DOCX"""
    
    def __init__(self):
        self.storage = DocumentStorage()
        self.validator = FieldValidator()
    
    def fill_document_from_path(self, template_path: str, fields: Dict[str, Any]) -> Document:
        """
        Preenche um documento a partir do caminho do template.
        Retorna o documento preenchido (Document) sem salvar.
        """
        # Carregar documento do template
        doc = Document(template_path)
        
        # Validar campos antes de preencher
        self.validator.validate_fields(fields)
        
        # Detectar tipo de comprador (PF ou PJ) e remover seção não utilizada
        buyer_type = self._detect_buyer_type(fields)
        if buyer_type:
            self._remove_unused_buyer_section(doc, buyer_type)
        
        # Preencher campos no documento
        self._replace_fields_in_document(doc, fields, field_mapping=None)
        
        # Corrigir textos verticais PRIMEIRO (antes de outras formatações)
        # Isso é importante para garantir que a tabela "VISTO DO COMPRADOR" seja corrigida
        self._fix_vertical_text(doc)
        
        # Ajustar formatação das linhas de assinatura
        self._format_signature_lines(doc)
        
        return doc
    
    async def fill_document(self, original_document_id: str, 
                           original_path: str, 
                           fields: Dict[str, Any],
                           field_mapping: Optional[Dict[str, str]] = None) -> str:
        """
        Preenche o documento com os dados fornecidos
        field_mapping: mapeamento field_id -> original_text para substituição precisa
        Retorna o ID do documento preenchido
        """
        # Carregar documento original
        doc = Document(original_path)
        
        # Validar campos antes de preencher
        self.validator.validate_fields(fields)
        
        # Preencher campos no documento
        self._replace_fields_in_document(doc, fields, field_mapping)
        
        # Salvar documento preenchido
        filled_document_id = f"{original_document_id}_filled"
        filled_path = self.storage.get_filled_file_path(filled_document_id)
        
        print(f"Salvando documento preenchido em: {filled_path}")
        doc.save(filled_path)
        
        # Verificar se o arquivo foi salvo corretamente
        import os
        if os.path.exists(filled_path):
            file_size = os.path.getsize(filled_path)
            print(f"Documento salvo com sucesso! Tamanho: {file_size} bytes")
        else:
            print(f"ERRO: Arquivo não foi salvo em {filled_path}")
        
        return filled_document_id
    
    def _replace_fields_in_document(self, doc: Document, 
                                    fields: Dict[str, Any],
                                    field_mapping: Optional[Dict[str, str]] = None):
        """
        Substitui placeholders {{CAMPO}} no documento pelos valores fornecidos
        """
        # Formatar valores antes de substituir
        formatted_fields = self._format_all_fields(fields)
        
        # Substituir em parágrafos
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, formatted_fields)
        
        # Substituir em tabelas (CUIDADO: não quebrar formatação)
        # Só processar células que contêm placeholders para evitar quebrar formatação de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Só substituir se houver placeholder
                        if '{{' in para.text:
                            self._replace_in_paragraph(para, formatted_fields)
    
    def _replace_in_paragraph(self, paragraph, fields: Dict[str, str]):
        """
        Substitui placeholders mantendo a formatação do parágrafo
        """
        # Combinar todos os runs em um texto único
        full_text = ''.join([run.text for run in paragraph.runs])
        
        # Verificar se há algum placeholder
        if '{{' not in full_text:
            return
        
        # Substituir todos os placeholders
        new_text = full_text
        for field_id, value in fields.items():
            placeholder = f'{{{{{field_id}}}}}'  # {{CAMPO}}
            new_text = new_text.replace(placeholder, str(value))
        
        # Se houve mudança, atualizar o parágrafo
        if new_text != full_text:
            # Preservar formatação do primeiro run
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.font.bold
                italic = first_run.font.italic
                
                # Limpar runs existentes
                for run in paragraph.runs:
                    run.text = ''
                
                # Adicionar novo texto com formatação preservada
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    if font_name:
                        paragraph.runs[0].font.name = font_name
                    if font_size:
                        paragraph.runs[0].font.size = font_size
                    paragraph.runs[0].font.bold = bold
                    paragraph.runs[0].font.italic = italic
            else:
                # Se não há runs, adicionar texto diretamente
                paragraph.text = new_text
    
    def _format_all_fields(self, fields: Dict[str, Any]) -> Dict[str, str]:
        """
        Formata todos os campos conforme seu tipo
        """
        formatted = {}
        
        for field_id, value in fields.items():
            if field_id in ROTA_DO_SOL_SCHEMA:
                field_def = ROTA_DO_SOL_SCHEMA[field_id]
                formatted[field_id] = self._format_value(value, field_def.type)
            else:
                formatted[field_id] = str(value) if value is not None else ''
        
        return formatted
    
    def _format_value(self, value: Any, field_type: FieldType) -> str:
        """
        Formata um valor conforme seu tipo
        """
        if value is None or value == '':
            return ''
        
        value_str = str(value)
        
        if field_type == FieldType.CPF:
            # Remove não-numéricos e formata
            digits = re.sub(r'\D', '', value_str)
            if len(digits) == 11:
                return f'{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}'
            return value_str
        
        elif field_type == FieldType.CNPJ:
            digits = re.sub(r'\D', '', value_str)
            if len(digits) == 14:
                return f'{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:]}'
            return value_str
        
        elif field_type == FieldType.CEP:
            digits = re.sub(r'\D', '', value_str)
            if len(digits) == 8:
                return f'{digits[:5]}-{digits[5:]}'
            return value_str
        
        elif field_type == FieldType.CURRENCY:
            # IMPORTANTE: NÃO adicionar "R$ " pois já está no template
            try:
                # Limpar valor e converter
                clean_value = re.sub(r'[^\d,.]', '', value_str)
                clean_value = clean_value.replace(',', '.')
                num = float(clean_value)
                # Formatar sem o prefixo R$
                formatted = f'{num:,.2f}'
                # Converter para formato brasileiro (1.234,56)
                formatted = formatted.replace(',', 'X').replace('.', ',').replace('X', '.')
                return formatted
            except:
                return value_str
        
        elif field_type == FieldType.PHONE:
            digits = re.sub(r'\D', '', value_str)
            if len(digits) == 9:
                return f'{digits[:5]}-{digits[5:]}'
            elif len(digits) == 8:
                return f'{digits[:4]}-{digits[4:]}'
            return value_str
        
        elif field_type == FieldType.DATE:
            # Converter de YYYY-MM-DD para DD/MM/YYYY se necessário
            if re.match(r'^\d{4}-\d{2}-\d{2}$', value_str):
                parts = value_str.split('-')
                return f'{parts[2]}/{parts[1]}/{parts[0]}'
            return value_str
        
        return value_str
    
    def _detect_buyer_type(self, fields: Dict[str, Any]) -> Optional[str]:
        """
        Detecta se o comprador é PF ou PJ baseado nos campos preenchidos.
        Retorna 'PF', 'PJ' ou None se não conseguir determinar.
        """
        has_pf_fields = any(key.startswith('COMPRADOR_PF_') for key in fields.keys() if fields.get(key))
        has_pj_fields = any(key.startswith('COMPRADOR_PJ_') for key in fields.keys() if fields.get(key))
        
        if has_pf_fields and not has_pj_fields:
            return 'PF'
        elif has_pj_fields and not has_pf_fields:
            return 'PJ'
        # Se ambos ou nenhum, retorna None (não remove nada)
        return None
    
    def _remove_unused_buyer_section(self, doc: Document, buyer_type: str):
        """
        Remove a seção do comprador não utilizada do documento.
        buyer_type: 'PF' ou 'PJ'
        IMPORTANTE: Não remove o preâmbulo ou outras seções do contrato.
        """
        # Determinar qual seção remover
        if buyer_type == 'PF':
            # Remover seção PJ
            section_prefix = 'COMPRADOR_PJ_'
            section_labels = ['Se for pessoa jurídica:']
        else:
            # Remover seção PF
            section_prefix = 'COMPRADOR_PF_'
            section_labels = ['Descrever os dados do comprador, se for pessoa física:']
        
        # Padrão regex para encontrar placeholders da seção não utilizada
        pattern = re.compile(r'\{\{' + re.escape(section_prefix) + r'[A-Z0-9_]+\}\}')
        
        # Identificar parágrafos a remover (usar índices para remoção segura)
        # IMPORTANTE: Só remover parágrafos que CONTÊM placeholders da seção não utilizada
        # Não remover parágrafos que apenas mencionam as palavras no contexto do preâmbulo
        paragraphs_to_remove = []
        
        # Encontrar onde começa a seção "1. COMPRADOR(ES):" para não remover nada antes
        comprador_section_start = -1
        for i, para in enumerate(doc.paragraphs):
            if '1. COMPRADOR' in para.text or 'COMPRADOR(ES):' in para.text:
                comprador_section_start = i
                break
        
        # Se não encontrou a seção, não remover nada (proteção)
        if comprador_section_start == -1:
            print("AVISO: Seção '1. COMPRADOR(ES):' não encontrada. Não removendo nada.")
            return
        
        # Só processar parágrafos APÓS o início da seção de compradores
        for i, para in enumerate(doc.paragraphs):
            # Pular preâmbulo e tudo antes da seção de compradores
            if i < comprador_section_start:
                continue
                
            para_text = para.text
            
            # Verificar se o parágrafo contém placeholders da seção não utilizada
            if pattern.search(para_text):
                paragraphs_to_remove.append(i)
            # Remover apenas labels específicos da seção (não palavras genéricas)
            elif any(label.lower() == para_text.strip().lower()[:len(label)] for label in section_labels):
                # Verificar se é realmente o label da seção (não apenas uma menção no texto)
                para_stripped = para_text.strip()
                if any(para_stripped.startswith(label) for label in section_labels):
                    paragraphs_to_remove.append(i)
        
        # Remover parágrafos de trás para frente (para não bagunçar índices)
        for i in reversed(paragraphs_to_remove):
            try:
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
            except:
                # Se não conseguir remover, pelo menos limpar o texto
                try:
                    doc.paragraphs[i].clear()
                except:
                    pass
        
        # Remover de tabelas também (mas com cuidado para não quebrar formatação)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_text = para.text
                        # Só remover se contém placeholders específicos
                        if pattern.search(para_text):
                            para.clear()
                        # Remover labels específicos
                        elif any(label.lower() == para_text.strip().lower()[:len(label)] for label in section_labels):
                            para_stripped = para_text.strip()
                            if any(para_stripped.startswith(label) for label in section_labels):
                                para.clear()
    
    def _format_signature_lines(self, doc: Document):
        """
        Ajusta o espaçamento das linhas de assinatura (Vendedora, Comprador, Testemunhas)
        para ter mais espaço para assinatura.
        """
        signature_keywords = ['Vendedora:', 'Vendedor:', 'Comprador:', 'Compradora:', 'Testemunhas:', 'Testemunha:']
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            
            # Verificar se é uma linha de assinatura
            is_signature_line = any(keyword.lower() in para_text.lower() for keyword in signature_keywords)
            
            if is_signature_line:
                # Aplicar espaçamento adequado
                para.paragraph_format.space_before = Pt(12)  # Espaço antes
                para.paragraph_format.space_after = Pt(24)   # Espaço depois (mais espaço para assinatura)
                
                # Se o parágrafo contém apenas o label (sem linha), adicionar espaço extra no próximo
                if not any(char in para_text for char in ['_', '─', '━', '─']):
                    # É apenas o label, garantir espaço no próximo parágrafo
                    para.paragraph_format.space_after = Pt(36)
        
        # Também ajustar em tabelas (caso as assinaturas estejam em tabelas)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_text = para.text.strip()
                        is_signature_line = any(keyword.lower() in para_text.lower() for keyword in signature_keywords)
                        
                        if is_signature_line:
                            para.paragraph_format.space_before = Pt(12)
                            para.paragraph_format.space_after = Pt(24)
                            
                            if not any(char in para_text for char in ['_', '─', '━', '─']):
                                para.paragraph_format.space_after = Pt(36)
    
    def _fix_vertical_text(self, doc: Document):
        """
        Corrige textos que estão na vertical, especialmente "VISTO DO COMPRADOR Ciente".
        Converte textos verticais para horizontais.
        IMPORTANTE: Processa TODAS as células de tabela que contêm texto vertical, mesmo sem placeholders.
        """
        visto_keywords = ['visto', 'comprador', 'ciente']
        
        # Processar parágrafos normais
        for para in doc.paragraphs:
            para_text = para.text.strip()
            
            # Verificar se contém texto de "visto" e está na vertical
            if any(keyword.lower() in para_text.lower() for keyword in visto_keywords):
                lines = para_text.split('\n')
                # Se tem múltiplas linhas e cada linha tem apenas 1-2 caracteres, provavelmente está vertical
                if len(lines) > 3 and all(len(line.strip()) <= 2 for line in lines if line.strip()):
                    # Converter para horizontal
                    horizontal_text = ' '.join(line.strip() for line in lines if line.strip())
                    
                    # Atualizar o parágrafo preservando formatação
                    if para.runs:
                        first_run = para.runs[0]
                        font_name = first_run.font.name
                        font_size = first_run.font.size
                        bold = first_run.font.bold
                        italic = first_run.font.italic
                        
                        for run in para.runs:
                            run.text = ''
                        
                        if para.runs:
                            para.runs[0].text = horizontal_text
                            if font_name:
                                para.runs[0].font.name = font_name
                            if font_size:
                                para.runs[0].font.size = font_size
                            para.runs[0].font.bold = bold
                            para.runs[0].font.italic = italic
                    else:
                        para.text = horizontal_text
                    
                    para.paragraph_format.alignment = None
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
        
        # Processar TODAS as tabelas - especialmente a tabela "VISTO DO COMPRADOR"
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Verificar se a célula contém texto relacionado a "visto do comprador"
                    cell_text = ' '.join([para.text.strip() for para in cell.paragraphs])
                    
                    if any(keyword.lower() in cell_text.lower() for keyword in visto_keywords):
                        # Processar todos os parágrafos da célula
                        for para_idx, para in enumerate(cell.paragraphs):
                            para_text = para.text.strip()
                            
                            if not para_text:
                                continue
                            
                            # Detectar se está na vertical:
                            # - Verificar se tem múltiplas linhas com 1-2 caracteres cada
                            lines = [line.strip() for line in para_text.split('\n') if line.strip()]
                            
                            is_vertical = False
                            
                            # Se tem mais de 3 linhas e a maioria tem 1-2 caracteres
                            if len(lines) > 3:
                                short_lines = sum(1 for line in lines if len(line) <= 2)
                                if short_lines >= len(lines) * 0.5:  # 50% das linhas são curtas
                                    is_vertical = True
                            # Se tem 2-3 linhas e todas são muito curtas (1-2 chars)
                            elif len(lines) >= 2:
                                if all(len(line) <= 2 for line in lines):
                                    is_vertical = True
                            
                            # Se está na vertical, converter para horizontal
                            if is_vertical:
                                # Juntar todas as linhas em uma única linha horizontal
                                horizontal_text = ' '.join(lines)
                                
                                print(f"Corrigindo texto vertical na tabela {table_idx}, linha {row_idx}, celula {cell_idx}:")
                                print(f"  Original (vertical): {repr(para_text[:100])}")
                                print(f"  Convertido (horizontal): {horizontal_text[:100]}")
                                
                                # Preservar formatação do primeiro run
                                if para.runs:
                                    first_run = para.runs[0]
                                    font_name = first_run.font.name
                                    font_size = first_run.font.size
                                    bold = first_run.font.bold
                                    italic = first_run.font.italic
                                    
                                    # Limpar todos os runs
                                    for run in para.runs:
                                        run.text = ''
                                    
                                    # Adicionar texto horizontal no primeiro run
                                    if para.runs:
                                        para.runs[0].text = horizontal_text
                                        if font_name:
                                            para.runs[0].font.name = font_name
                                        if font_size:
                                            para.runs[0].font.size = font_size
                                        para.runs[0].font.bold = bold
                                        para.runs[0].font.italic = italic
                                else:
                                    # Se não tem runs, criar novo parágrafo
                                    para.text = horizontal_text
                                
                                # Ajustar alinhamento para horizontal
                                para.paragraph_format.alignment = None
                                
                                # Tentar ajustar orientação da célula (se possível)
                                try:
                                    # Tentar forçar orientação horizontal na célula
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                                except:
                                    pass
    
