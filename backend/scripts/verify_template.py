"""
Script para verificar e corrigir o template DOCX
"""
import sys
from pathlib import Path

# Adicionar o diretório raiz ao path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def verify_and_fix_template():
    """Verifica e corrige o template DOCX"""
    template_path = Path(__file__).parent.parent / "templates" / "CONTRATO_ROTA_DO_SOL_TEMPLATE.docx"
    
    if not template_path.exists():
        print(f"ERRO: Template não encontrado em {template_path}")
        return False
    
    print(f"Carregando template: {template_path}")
    doc = Document(str(template_path))
    
    # Extrair todo o texto
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    
    full_text = '\n'.join(all_text)
    
    # Verificar se o preâmbulo está presente
    preambulo_keywords = [
        "CONTRATO PARTICULAR DE PROMESSA DE COMPRA E VENDA",
        "Quadro Resumo",
        "LALU – ADMINISTRADORA DE BENS LTDA",
        "RESIDENCIAL ROTA DO SOL"
    ]
    
    missing_keywords = []
    for keyword in preambulo_keywords:
        if keyword not in full_text:
            missing_keywords.append(keyword)
    
    if missing_keywords:
        print(f"AVISO: Palavras-chave do preambulo nao encontradas: {missing_keywords}")
        print("O preambulo pode estar faltando ou incompleto.")
    else:
        print("OK: Preambulo encontrado no template")
    
    # Verificar seção de compradores
    if "1. COMPRADOR(ES):" in full_text:
        print("OK: Secao '1. COMPRADOR(ES):' encontrada")
    else:
        print("ERRO: Secao '1. COMPRADOR(ES):' NAO encontrada")
    
    # Verificar tabela "VISTO DO COMPRADOR"
    visto_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = ' '.join([para.text for para in cell.paragraphs])
                if 'visto' in cell_text.lower() and 'comprador' in cell_text.lower():
                    visto_found = True
                    print(f"OK: Tabela 'VISTO DO COMPRADOR' encontrada")
                    print(f"  Texto da celula: {cell_text[:50]}...")
                    break
            if visto_found:
                break
        if visto_found:
            break
    
    if not visto_found:
        print("ERRO: Tabela 'VISTO DO COMPRADOR' NAO encontrada")
    
    # Listar primeiros parágrafos para debug
    print("\nPrimeiros 10 parágrafos do documento:")
    for i, para in enumerate(doc.paragraphs[:10]):
        text = para.text.strip()
        if text:
            print(f"  {i+1}. {text[:80]}...")
    
    return True

if __name__ == "__main__":
    verify_and_fix_template()
